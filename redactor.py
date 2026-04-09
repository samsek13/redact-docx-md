"""
文件脱敏与恢复核心模块
支持格式：.txt, .docx
"""

import re
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Tuple, Dict, List

# 配置日志
logger = logging.getLogger(__name__)

# 尝试导入 python-docx
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class Redactor:
    """脱敏器：负责文本脱敏和恢复"""

    def __init__(self):
        self.counter = 0
        self.mapping: Dict[str, str] = {}

    def reset(self):
        """重置计数器和映射"""
        self.counter = 0
        self.mapping = {}

    def redact_text(self, text: str, words: List[str]) -> str:
        """
        对文本进行脱敏处理

        Args:
            text: 原始文本
            words: 需要脱敏的词列表

        Returns:
            脱敏后的文本
        """
        if not text or not words:
            return text

        # 过滤空词并去重
        words = list(set(w for w in words if w and w.strip()))
        if not words:
            return text

        # 按词长降序排序，避免短词误匹配
        words_sorted = sorted(words, key=len, reverse=True)

        def replace_match(match):
            """替换匹配到的词"""
            matched_word = match.group(0)
            # 检查是否已有映射（同一词可能多次出现）
            for placeholder, original in self.mapping.items():
                if original == matched_word:
                    return placeholder

            # 创建新映射
            self.counter += 1
            placeholder = f"[REDACTED_{self.counter}]"
            self.mapping[placeholder] = matched_word
            return placeholder

        # 为每个词创建正则模式并替换
        result = text
        for word in words_sorted:
            # 使用 re.escape 处理特殊字符，确保精确匹配
            pattern = re.escape(word)
            result = re.sub(pattern, replace_match, result, flags=re.IGNORECASE)

        return result

    def deredact_text(self, text: str, mapping: Dict[str, str]) -> str:
        """
        根据映射字典恢复脱敏文本

        Args:
            text: 脱敏后的文本
            mapping: 占位符到原文的映射字典

        Returns:
            恢复后的文本
        """
        if not text or not mapping:
            return text

        result = text
        # 按占位符长度降序排序，避免部分匹配问题
        sorted_mappings = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)

        for placeholder, original in sorted_mappings:
            result = result.replace(placeholder, original)

        return result


def redact_txt(file_path: Path, words: List[str], output_dir: Path = None) -> Tuple[Path, Path]:
    """
    处理 txt 文件脱敏

    Args:
        file_path: 原始 txt 文件路径
        words: 需要脱敏的词列表
        output_dir: 输出目录，默认为原文件所在目录

    Returns:
        (脱敏文件路径, 映射文件路径)
    """
    if output_dir is None:
        output_dir = file_path.parent

    # 读取文件
    text = file_path.read_text(encoding='utf-8')

    # 脱敏处理
    redactor = Redactor()
    redacted_text = redactor.redact_text(text, words)

    # 生成输出文件名
    stem = file_path.stem
    output_file = output_dir / f"{stem}_redacted.txt"
    mapping_file = output_dir / f"{stem}_redacted.mapping.json"

    # 写入脱敏文件
    output_file.write_text(redacted_text, encoding='utf-8')

    # 写入映射文件
    mapping_data = {
        "original_filename": file_path.name,
        "created_at": datetime.now().isoformat(),
        "mappings": redactor.mapping
    }
    mapping_file.write_text(json.dumps(mapping_data, ensure_ascii=False, indent=2), encoding='utf-8')

    return output_file, mapping_file


def deredact_txt(file_path: Path, mapping_file: Path, output_dir: Path = None) -> Path:
    """
    恢复 txt 脱敏文件

    Args:
        file_path: 脱敏 txt 文件路径
        mapping_file: 映射文件路径
        output_dir: 输出目录，默认为原文件所在目录

    Returns:
        恢复后的文件路径
    """
    if output_dir is None:
        output_dir = file_path.parent

    # 读取文件
    text = file_path.read_text(encoding='utf-8')
    mapping_data = json.loads(mapping_file.read_text(encoding='utf-8'))

    # 恢复处理
    redactor = Redactor()
    restored_text = redactor.deredact_text(text, mapping_data['mappings'])

    # 生成输出文件名
    stem = file_path.stem.replace('_redacted', '')
    output_file = output_dir / f"{stem}_restored.txt"

    # 写入恢复文件
    output_file.write_text(restored_text, encoding='utf-8')

    return output_file


def _redact_hyperlink_urls(redactor, paragraph, words):
    """脱敏超链接的 URL（存储在 relationship 中）"""
    from docx.oxml.ns import qn

    for hyperlink in paragraph._element.findall(qn('w:hyperlink')):
        rId = hyperlink.get(qn('r:id'))
        if not rId:
            continue
        try:
            rel = paragraph.part.rels[rId]
            url = str(rel.target_ref)
            redacted_url = redactor.redact_text(url, words)
            if redacted_url != url:
                # 注意：python-docx 无官方 API 修改超链接 URL，
                # 需直接修改内部属性 _target（存在版本兼容风险）
                rel._target = redacted_url
                logger.debug(f"超链接 URL 已脱敏: {url} -> {redacted_url}")
        except KeyError:
            logger.debug(f"跳过无效超链接关系 ID: {rId}")


def _redact_comments(redactor, doc, words):
    """脱敏 word/comments.xml 等附加 XML 部件中的文本"""
    from docx.oxml.ns import qn

    for part in doc.part.package.iter_parts():
        partname = str(part.partname)
        # 处理 comments.xml, footnotes.xml, endnotes.xml 等（排除 document.xml）
        if '/word/' in partname and partname.endswith('.xml') and 'document.xml' not in partname:
            if not hasattr(part, '_element'):
                continue
            try:
                # 直接修改 part._element（blob 会从 _element 序列化）
                for t_elem in part._element.iter(qn('w:t')):
                    if t_elem.text:
                        redacted = redactor.redact_text(t_elem.text, words)
                        if redacted != t_elem.text:
                            t_elem.text = redacted
                for t_elem in part._element.iter(qn('w:delText')):
                    if t_elem.text:
                        redacted = redactor.redact_text(t_elem.text, words)
                        if redacted != t_elem.text:
                            t_elem.text = redacted
            except Exception as e:
                logger.debug(f"跳过无法处理的 XML 部件 {partname}: {e}")


def _redact_cross_element_text(elements, redactor, words):
    """处理跨元素拆分的文本脱敏

    将多个相邻文本元素的内容拼接后整体脱敏，解决敏感词被拆分到不同元素的问题。

    Args:
        elements: 文本元素列表（如 w:t 或 w:delText）
        redactor: 脱敏器实例
        words: 敏感词列表
    """
    all_text = ''.join(e.text for e in elements if e.text)
    redacted = redactor.redact_text(all_text, words)
    if redacted != all_text and elements:
        elements[0].text = redacted
        for e in elements[1:]:
            e.text = ""


def _redact_textbox_content(doc, redactor, words):
    """脱敏文本框内容（w:txbxContent 内的段落）

    文本框内的文本存储在独立的 w:txbxContent 元素中，需单独处理。
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    for textbox in doc.element.body.iter(qn('w:txbxContent')):
        for p_elem in textbox.iter(qn('w:p')):
            paragraph = Paragraph(p_elem, doc)
            _redact_paragraph(redactor, paragraph, words)


def _deredact_textbox_content(doc, redactor, mapping):
    """恢复文本框内容"""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    for textbox in doc.element.body.iter(qn('w:txbxContent')):
        for p_elem in textbox.iter(qn('w:p')):
            paragraph = Paragraph(p_elem, doc)
            _deredact_paragraph(redactor, paragraph, mapping)


def _deredact_comments(redactor, doc, mapping):
    """恢复 word/comments.xml 等附加 XML 部件中的文本"""
    from docx.oxml.ns import qn

    for part in doc.part.package.iter_parts():
        partname = str(part.partname)
        if '/word/' in partname and partname.endswith('.xml') and 'document.xml' not in partname:
            if not hasattr(part, '_element'):
                continue
            try:
                for t_elem in part._element.iter(qn('w:t')):
                    if t_elem.text:
                        restored = redactor.deredact_text(t_elem.text, mapping)
                        if restored != t_elem.text:
                            t_elem.text = restored
                for t_elem in part._element.iter(qn('w:delText')):
                    if t_elem.text:
                        restored = redactor.deredact_text(t_elem.text, mapping)
                        if restored != t_elem.text:
                            t_elem.text = restored
            except Exception:
                pass


def _redact_paragraph(redactor, paragraph, words):
    """多阶段脱敏：处理段落内所有文本，包括跨元素拆分"""
    from docx.oxml.ns import qn

    # Phase 1: 逐 run 处理（保留格式）
    for run in paragraph.runs:
        if run.text:
            run.text = redactor.redact_text(run.text, words)

    # Phase 2: 段落级兜底（捕获跨 run 拆分词）
    remaining = ''.join(run.text or '' for run in paragraph.runs)
    redacted_remaining = redactor.redact_text(remaining, words)
    if redacted_remaining != remaining:
        if paragraph.runs:
            paragraph.runs[0].text = redacted_remaining
            for run in paragraph.runs[1:]:
                run.text = ""

    # 处理段落内所有 w:t（覆盖超链接、w:ins、w:sdt 等）
    for t_elem in paragraph._element.iter(qn('w:t')):
        if t_elem.text:
            redacted = redactor.redact_text(t_elem.text, words)
            if redacted != t_elem.text:
                t_elem.text = redacted

    # 处理所有 w:delText（修订追踪中的删除文本）
    for t_elem in paragraph._element.iter(qn('w:delText')):
        if t_elem.text:
            redacted = redactor.redact_text(t_elem.text, words)
            if redacted != t_elem.text:
                t_elem.text = redacted

    # Phase 3: 全段落拼接兜底（捕获跨元素拆分的敏感词）
    # 使用辅助函数处理 w:t 和 w:delText 的跨元素拆分
    all_t_elems = [t for t in paragraph._element.iter(qn('w:t')) if t.text]
    _redact_cross_element_text(all_t_elems, redactor, words)

    all_dt_elems = [t for t in paragraph._element.iter(qn('w:delText')) if t.text]
    _redact_cross_element_text(all_dt_elems, redactor, words)

    # 处理超链接 URL
    _redact_hyperlink_urls(redactor, paragraph, words)


def _deredact_paragraph(redactor, paragraph, mapping):
    """对段落进行 run 级别恢复，包括所有文本元素和 URL"""
    from docx.oxml.ns import qn

    # 恢复所有 w:t 元素（覆盖普通 runs、超链接、w:ins、w:sdt 等）
    for t_elem in paragraph._element.iter(qn('w:t')):
        if t_elem.text:
            restored = redactor.deredact_text(t_elem.text, mapping)
            if restored != t_elem.text:
                t_elem.text = restored

    # 恢复所有 w:delText 元素
    for t_elem in paragraph._element.iter(qn('w:delText')):
        if t_elem.text:
            restored = redactor.deredact_text(t_elem.text, mapping)
            if restored != t_elem.text:
                t_elem.text = restored

    # 恢复超链接 URL
    for hyperlink in paragraph._element.findall(qn('w:hyperlink')):
        rId = hyperlink.get(qn('r:id'))
        if not rId:
            continue
        try:
            rel = paragraph.part.rels[rId]
            url = str(rel.target_ref)
            restored_url = redactor.deredact_text(url, mapping)
            if restored_url != url:
                # 注意：python-docx 无官方 API 修改超链接 URL，
                # 需直接修改内部属性 _target（存在版本兼容风险）
                rel._target = restored_url
        except KeyError:
            logger.debug(f"跳过无效超链接关系 ID: {rId}")


def _process_all_paragraphs(doc, process_fn):
    """遍历文档中所有段落（正文、所有层级表格、页眉页脚）并执行处理函数

    使用 XML 级迭代确保覆盖所有位置，包括：
    - 嵌套表格、合并单元格
    - w:sdt（结构化文档标签）内的段落
    - 修订追踪（w:ins/w:del）内的文本
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    # 正文及所有表格（含嵌套）：遍历 body 中全部 w:p 元素
    for p_elem in doc.element.body.iter(qn('w:p')):
        paragraph = Paragraph(p_elem, doc)
        process_fn(paragraph)

    # 页眉页脚
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_fn(paragraph)
        for paragraph in section.footer.paragraphs:
            process_fn(paragraph)

def redact_docx(file_path: Path, words: List[str], output_dir: Path = None) -> Tuple[Path, Path]:
    """
    处理 docx 文件脱敏

    Args:
        file_path: 原始 docx 文件路径
        words: 需要脱敏的词列表
        output_dir: 输出目录，默认为原文件所在目录

    Returns:
        (脱敏文件路径, 映射文件路径)
    """
    if not HAS_DOCX:
        raise ImportError("python-docx 未安装，无法处理 docx 文件。请运行: pip install python-docx")

    if output_dir is None:
        output_dir = file_path.parent

    # 打开文档
    doc = Document(str(file_path))
    redactor = Redactor()

    # 遍历所有段落进行 run 级别脱敏（包括超链接）
    _process_all_paragraphs(doc, lambda p: _redact_paragraph(redactor, p, words))

    # 脱敏文本框内容
    _redact_textbox_content(doc, redactor, words)

    # 脱敏批注、脚注等附加 XML 部件
    _redact_comments(redactor, doc, words)

    # 生成输出文件名
    stem = file_path.stem
    output_file = output_dir / f"{stem}_redacted.docx"
    mapping_file = output_dir / f"{stem}_redacted.mapping.json"

    # 保存文档
    doc.save(str(output_file))

    # 写入映射文件
    mapping_data = {
        "original_filename": file_path.name,
        "created_at": datetime.now().isoformat(),
        "mappings": redactor.mapping
    }
    mapping_file.write_text(json.dumps(mapping_data, ensure_ascii=False, indent=2), encoding='utf-8')

    return output_file, mapping_file


def deredact_docx(file_path: Path, mapping_file: Path, output_dir: Path = None) -> Path:
    """
    恢复 docx 脱敏文件

    Args:
        file_path: 脱敏 docx 文件路径
        mapping_file: 映射文件路径
        output_dir: 输出目录，默认为原文件所在目录

    Returns:
        恢复后的文件路径
    """
    if not HAS_DOCX:
        raise ImportError("python-docx 未安装，无法处理 docx 文件。请运行: pip install python-docx")

    if output_dir is None:
        output_dir = file_path.parent

    # 打开文档
    doc = Document(str(file_path))
    mapping_data = json.loads(mapping_file.read_text(encoding='utf-8'))
    mapping = mapping_data['mappings']
    redactor = Redactor()

    # 遍历所有段落进行 run 级别恢复（包括超链接）
    _process_all_paragraphs(doc, lambda p: _deredact_paragraph(redactor, p, mapping))

    # 恢复文本框内容
    _deredact_textbox_content(doc, redactor, mapping)

    # 恢复批注、脚注等附加 XML 部件
    _deredact_comments(redactor, doc, mapping)

    # 生成输出文件名
    stem = file_path.stem.replace('_redacted', '')
    output_file = output_dir / f"{stem}_restored.docx"

    # 保存文档
    doc.save(str(output_file))

    return output_file


def get_supported_extensions() -> List[str]:
    """获取支持的文件扩展名"""
    extensions = ['.txt', '.md']
    if HAS_DOCX:
        extensions.append('.docx')
    return extensions


def is_supported(file_path: Path) -> bool:
    """检查文件是否支持"""
    return file_path.suffix.lower() in get_supported_extensions()